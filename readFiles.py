from docx import Document
from io import BytesIO
import re

def _open_docx(file_or_path):
    if hasattr(file_or_path, "read"):
        data = file_or_path.read()
        try:
            file_or_path.seek(0)
        except Exception:
            pass
        return Document(BytesIO(data))
    elif isinstance(file_or_path, str):
        return Document(file_or_path)
    else:
        raise ValueError("Unsupported input type.")

def _iter_block_items_correct(doc):
    body_elm = doc.element.body
    paragraphs = doc.paragraphs
    tables = doc.tables
    p_idx, t_idx = 0, 0
    for child in body_elm.iterchildren():
        tag = child.tag
        if tag.endswith('}p'):
            yield ('p', paragraphs[p_idx])
            p_idx += 1
        elif tag.endswith('}tbl'):
            yield ('tbl', tables[t_idx])
            t_idx += 1

def extract_number_heading(file_or_path, debug=False):
    """
    Extracts DOCX sections correctly using Heading 1/2, while keeping
    numbered/bulleted paragraphs as content (not new sections).
    """
    doc = _open_docx(file_or_path)
    sections = []
    parent_heading = None
    current_section = None
    current_content = []
    auto_idx = 0

    def flush_section():
        nonlocal current_section, current_content
        if current_section:
            text = "\n".join(current_content).strip()
            sections.append((current_section[0], current_section[1], text))
        current_content = []

    def start_new_section(heading_text):
        nonlocal auto_idx, current_section, current_content
        flush_section()
        auto_idx += 1
        current_section = (str(auto_idx), heading_text)
        current_content = []

    numbered_line_pattern = re.compile(r"^\s*\d+[\.\)]\s+")
    bullet_line_pattern = re.compile(r"^\s*[\u2022\-\*]\s+")
    short_num_pattern = re.compile(r"^\s*\d+(\.\d+)*\s*[A-Za-z]")

    for btype, block in _iter_block_items_correct(doc):
        if btype == 'p':
            para = block
            text = (para.text or "").strip()
            if not text:
                continue

            try:
                style_name = (para.style.name or "").lower()
            except Exception:
                style_name = ""

            # --- Heading 1 (main section)
            if style_name.startswith("heading 1"):
                flush_section()
                parent_heading = text
                continue

            # --- Heading 2 (subsection), but filter false positives ---
            if style_name.startswith("heading 2"):
                # reject numbered or bulleted lines as real sections
                if numbered_line_pattern.match(text) or bullet_line_pattern.match(text):
                    # just append as content under current section
                    if not current_section:
                        start_new_section("Preamble")
                    current_content.append(text)
                    continue

                # real heading 2
                heading_text = f"{parent_heading} :: {text}" if parent_heading else text
                start_new_section(heading_text)
                continue

            # --- Lists / Normal paragraphs ---
            if (style_name.startswith("list") or
                numbered_line_pattern.match(text) or
                bullet_line_pattern.match(text)):
                if not current_section:
                    start_new_section("Preamble")
                current_content.append(text)
                continue

            # --- Bold short text (maybe heading-like) ---
            is_bold = any(run.bold for run in para.runs) if para.runs else False
            if is_bold and len(text) < 80 and not style_name.startswith("list"):
                # treat as inline sub-heading but part of same section
                current_content.append(f"**{text}**")
                continue

            # --- Default: content ---
            if not current_section:
                start_new_section("Preamble")
            current_content.append(text)

        elif btype == 'tbl':
            table = block
            rows_text = []
            for row in table.rows:
                row_text = " | ".join(
                    [" / ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()]) for cell in row.cells]
                )
                if row_text:
                    rows_text.append(row_text)
            if rows_text:
                if not current_section:
                    start_new_section("Preamble")
                current_content.append("TABLE_START")
                current_content.extend(rows_text)
                current_content.append("TABLE_END")

    flush_section()

    if debug:
        print(f"Extracted {len(sections)} sections")
        for n, h, c in sections:
            print(f"[{n}] {h} -> {len(c)} chars")

    return sections
