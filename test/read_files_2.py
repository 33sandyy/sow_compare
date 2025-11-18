import os
import re
import json
from docx import Document

# readFiles.py
import re
from docx import Document
from io import BytesIO

# readFiles.py
import os
import re
from io import BytesIO
from docx import Document

def _open_docx(file_or_path):
    """
    Handles both a file path or a file-like object (e.g., Flask upload or BytesIO).
    """
    if hasattr(file_or_path, "read"):
        # Flask uploaded file or BytesIO
        data = file_or_path.read()
        file_or_path.seek(0)
        return Document(BytesIO(data))
    elif isinstance(file_or_path, str):
        return Document(file_or_path)
    else:
        raise ValueError("Unsupported input type")

from docx import Document
import re

from docx import Document
import re

def extract_number_heading(docx_path):
    """
    Extracts headings and content (including table text) from DOCX.
    Works with numbered or plain headings like '1. Scope of Work' or 'Scope of Work'.
    """
    doc = Document(docx_path)
    sections = []
    current_section = None
    current_content = []

    # Accept numbered or plain uppercase / bold headings
    numbered_pattern = re.compile(r"^(\d+(\.\d+)*)\s*(.*)$")

    def add_section():
        nonlocal current_section, current_content
        if current_section:
            sections.append((current_section[0], current_section[1], "\n".join(current_content).strip()))
            current_content = []

    def read_paragraphs(paragraphs):
        texts = []
        for para in paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)
        return texts

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        m = numbered_pattern.match(text)
        if m and m.group(3):
            # “1. Scope of Work” → ('1', 'Scope of Work')
            add_section()
            sec_num = m.group(1)
            heading = m.group(3)
            current_section = (sec_num, heading)
        elif para.style.name.startswith("Heading"):
            # Word heading style (no numbers)
            add_section()
            sec_num = str(len(sections) + 1)
            current_section = (sec_num, text)
        else:
            current_content.append(text)

    # --- include tables ---
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            cells = [" | ".join(read_paragraphs(cell.paragraphs)) for cell in row.cells]
            row_text = " | ".join([c for c in cells if c])
            if row_text:
                table_text.append(row_text)
        if table_text and current_section:
            current_content.append("TABLE DATA:\n" + "\n".join(table_text))

    add_section()
    return sections

if __name__ == "__main__":
    # Example usage:
    sow_file = "static/SOW_SDD2_SOW for Full Stack Developer DX Actvities11-Feb-2025.docx"
    template_file = "static/SOWTemplateDoc.docx"

    sections = extract_number_heading(sow_file, debug=True)

    for sec in sections:
        print(f"\n[{sec[0]}] {sec[1]}")
        print(sec[2][:200], "...")