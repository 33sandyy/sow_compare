import os
import re
import json
from docx import Document

# readFiles.py
import re
from docx import Document
from io import BytesIO

# readFiles.py
import os
import re
from io import BytesIO
from docx import Document

# readFiles.py (replace or add this function)
from docx import Document
from io import BytesIO
import re

def _open_docx(file_or_path):
    """Return Document object from path or file-like object."""
    if hasattr(file_or_path, "read"):
        data = file_or_path.read()
        try:
            file_or_path.seek(0)
        except Exception:
            pass
        return Document(BytesIO(data))
    elif isinstance(file_or_path, str):
        return Document(file_or_path)
    else:
        raise ValueError("Unsupported input type. Provide path or file-like object.")

def _iter_block_items(parent):
    """
    Yield block-level items (paragraphs and tables) in document order.
    Each yielded item is a tuple ('p', paragraph) or ('tbl', table).
    """
    # Based on python-docx internal XML
    for child in parent.element.body.iterchildren():
        if child.tag.endswith('}p'):   # paragraph
            yield ('p', parent.paragraphs[parent.element.body.index(child)])
        elif child.tag.endswith('}tbl'):  # table
            # find matching table by xml index
            # parent.tables are in document order, but we need to map properly
            # We'll yield tables sequentially from parent.tables in order encountered
            # So we track an index externally in caller (simpler). For safety, fallback:
            for tbl in parent.tables:
                # yield the first table we haven't yielded — but to keep simple, use xml child
                yield ('tbl', tbl)
                break

def _iter_block_items_correct(doc):
    """
    Reliable iterator (works across python-docx versions).
    This inspects the XML children of the document body and yields in order.
    """
    body_elm = doc.element.body
    paragraphs = doc.paragraphs
    tables = doc.tables
    p_idx = 0
    t_idx = 0

    for child in body_elm:
        tag = child.tag
        if tag.endswith('}p'):
            # next paragraph
            yield ('p', paragraphs[p_idx])
            p_idx += 1
        elif tag.endswith('}tbl'):
            yield ('tbl', tables[t_idx])
            t_idx += 1
        else:
            # unrecognized block — skip
            continue

def extract_number_heading(file_or_path, debug=False):
    """
    Extract headings (numbered or heading styles) and their content,
    preserving document order and including table contents in the correct section.

    Returns: list of tuples -> (section_number, section_heading, section_content)
    """
    doc = _open_docx(file_or_path)

    sections = []
    current_section = None    # tuple (sec_num, heading)
    current_content_lines = []

    # pattern to catch numbered headings like "1.", "1.2", "2 - Heading"
    numbered_pattern = re.compile(r"^\s*(\d+(?:[\.\-]\d+)*)[\.\-\)]?\s*(.+)$")

    auto_idx = 0

    def flush_section():
        nonlocal current_section, current_content_lines
        if current_section:
            text = "\n".join(current_content_lines).strip()
            sections.append((current_section[0], current_section[1], text))
            current_content_lines = []

    # iterate blocks in document order
    for btype, block in _iter_block_items_correct(doc):
        if btype == 'p':
            para = block
            text = (para.text or "").strip()
            if not text:
                # empty paragraph — skip
                continue

            # safe style name access
            try:
                style_name = (para.style.name or "").strip()
            except Exception:
                style_name = ""

            # 1) numbered heading like "1. Scope"
            m = numbered_pattern.match(text)
            if m:
                # start new section
                flush_section()
                sec_num = m.group(1).strip()
                heading = m.group(2).strip()
                current_section = (sec_num, heading)
                continue

            # 2) Word heading style (Heading 1/2/3, case-insensitive)
            if style_name and style_name.lower().startswith("heading"):
                flush_section()
                auto_idx += 1
                current_section = (str(auto_idx), text)
                continue

            # 3) short bold paragraph heuristic (possible heading)
            is_bold = False
            try:
                is_bold = any(run.bold for run in para.runs) if para.runs else False
            except Exception:
                is_bold = False

            if is_bold and len(text) < 160:
                flush_section()
                auto_idx += 1
                current_section = (str(auto_idx), text)
                continue

            # otherwise treat as content — if no current_section yet, create a preamble
            if not current_section:
                auto_idx += 1
                current_section = (str(auto_idx), "Preamble")
            current_content_lines.append(text)

        elif btype == 'tbl':
            table = block
            # Read rows, preserve structure: join cells with ' | '
            rows_text = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    # collect paragraph text inside cell
                    cell_paras = [p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()]
                    if cell_paras:
                        # join paragraphs inside a cell with ' / ' or ' '
                        cells.append(" / ".join(cell_paras))
                    else:
                        cells.append("")  # keep place
                # join cells with ' | ' to preserve column boundaries
                row_text = " | ".join([c for c in cells if c])
                if row_text:
                    rows_text.append(row_text)

            if rows_text:
                if not current_section:
                    auto_idx += 1
                    current_section = (str(auto_idx), "Preamble")
                # mark table start and then append rows
                current_content_lines.append("TABLE_START")
                current_content_lines.extend(rows_text)
                current_content_lines.append("TABLE_END")

    # flush last section
    flush_section()

    if debug:
        print(f"Extracted {len(sections)} sections.")
        for sec in sections:
            sn, sh, sc = sec
            print(f"[{sn}] {sh} → {len(sc)} chars")

    return sections


if __name__ == "__main__":
    # Example usage:
    sow_file = "static/SOW_SDD2_SOW for Full Stack Developer DX Actvities11-Feb-2025.docx"
    template_file = "static/SOWTemplateDoc.docx"

    sections = extract_number_heading(sow_file, debug=True)

    for sec in sections:
        print(f"\n[{sec[0]}] {sec[1]}")
        print(sec[2][:200], "...")