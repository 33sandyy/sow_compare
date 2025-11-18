import os
import re
import json
from docx import Document

# readFiles.py
import re
from docx import Document
from io import BytesIO

# readFiles.py
import os
import re
from io import BytesIO
from docx import Document

def _open_docx(file_or_path):
    """
    Handles both a file path or a file-like object (e.g., Flask upload or BytesIO).
    """
    if hasattr(file_or_path, "read"):
        # Flask uploaded file or BytesIO
        data = file_or_path.read()
        file_or_path.seek(0)
        return Document(BytesIO(data))
    elif isinstance(file_or_path, str):
        return Document(file_or_path)
    else:
        raise ValueError("Unsupported input type")

def extract_number_heading(file_or_path, debug=False):
    """
    Extracts all Heading 2 sections and their content from a .docx file.
    Returns a list of tuples: (section_number, section_heading, section_content)
    """

    doc = _open_docx(file_or_path)

    sections = []
    current_section = None
    section_counter = 0
    style_summary = {}

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Get style name safely
        try:
            style_name = (para.style.name or "").strip().lower()
        except Exception:
            style_name = ""

        style_summary[style_name] = style_summary.get(style_name, 0) + 1

        # Detect Heading 2 (primary section headers)
        if "heading 2" in style_name:
            if current_section:
                sections.append(current_section)

            section_counter += 1
            current_section = {
                "section_number": str(section_counter),
                "section_heading": text,
                "section_content": ""
            }

        # Otherwise, treat as content
        else:
            if current_section:
                current_section["section_content"] += text + "\n"

    # Append the last section
    if current_section:
        sections.append(current_section)

    if debug:
        print("\n=== Style Summary ===")
        for style, count in sorted(style_summary.items(), key=lambda x: -x[1]):
            print(f"{style}: {count}")
        print("=====================\n")
        print(f"Extracted {len(sections)} sections.\n")

    return [(s["section_number"], s["section_heading"], s["section_content"].strip()) for s in sections]




if __name__ == "__main__":
    # Example usage:
    sow_file = "static/SOW_SDD2_SOW for Full Stack Developer DX Actvities11-Feb-2025.docx"
    template_file = "static/SOWTemplateDoc.docx"

    sections = extract_number_heading(sow_file, debug=True)

    for sec in sections:
        print(f"\n[{sec[0]}] {sec[1]}")
        print(sec[2][:200], "...")